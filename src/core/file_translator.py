<<<<<<< HEAD
# src/core/file_translator.py
"""
Professional Multi-Format File Translation Engine (Optimized)
Supports: .docx, .bdf, .jpg/.png/.jpeg, .pdf
Uses: deep_translator (free, no API key)
Architecture: Extensible handler pattern for future formats (pptx, xlsx)
Author: Alsaeed Alwazzan

Performance Optimizations:
- Batch translation: combine texts with separator → single API call → split back
- Larger chunks (4500 chars) to minimize API roundtrips
- Lazy imports to avoid startup delay
- OCR is optional (easyocr or pytesseract, with graceful fallback)
"""

import io
import os
import re
import time
import logging
from datetime import datetime
from typing import Callable, Optional, Dict, List, Tuple, Any

# ──────────────────────────────────────────────
# Constants
# ──────────────────────────────────────────────
BATCH_SEPARATOR = "\n|||SPLIT|||\n"
CHUNK_SIZE = 4500  # Max chars per translation request


# ──────────────────────────────────────────────
# Translation Service (Free, No API Key)
# ──────────────────────────────────────────────

class TranslationService:
    """Fast, chunked translation using deep_translator (free Google Translate)."""

    SUPPORTED_LANGUAGES = {
        "ar": "العربية (Arabic)",
        "en": "English",
        "fr": "الفرنسية (French)",
        "es": "الإسبانية (Spanish)",
        "de": "الألمانية (German)",
        "tr": "التركية (Turkish)",
        "ur": "الأردية (Urdu)",
        "hi": "الهندية (Hindi)",
        "zh-CN": "الصينية (Chinese)",
        "ja": "اليابانية (Japanese)",
        "ko": "الكورية (Korean)",
        "ru": "الروسية (Russian)",
        "pt": "البرتغالية (Portuguese)",
        "it": "الإيطالية (Italian)",
        "id": "الإندونيسية (Indonesian)",
        "ms": "الملايوية (Malay)",
        "tl": "الفلبينية (Filipino)",
        "bn": "البنغالية (Bengali)",
        "th": "التايلاندية (Thai)",
        "vi": "الفيتنامية (Vietnamese)",
    }

    def __init__(self, source_lang: str = "auto", target_lang: str = "ar", logger=None):
        from deep_translator import GoogleTranslator
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.logger = logger
        self._translator = GoogleTranslator(source=source_lang, target=target_lang)

    def translate_text(self, text: str, progress_callback: Optional[Callable] = None) -> str:
        """Translate a single text string, with chunking for large content."""
        if not text or not text.strip():
            return ""

        # If small enough, translate directly
        if len(text) <= CHUNK_SIZE:
            try:
                result = self._translator.translate(text)
                if progress_callback:
                    progress_callback(1.0, "✅")
                return result if result else text
            except Exception as e:
                logging.warning(f"Translation failed: {e}")
                return text

        # Split large text into chunks at newline boundaries
        chunks = self._split_into_chunks(text)
        total = len(chunks)
        translated_parts = []

        for i, chunk in enumerate(chunks):
            try:
                result = self._translator.translate(chunk)
                translated_parts.append(result if result else chunk)
            except Exception as e:
                logging.warning(f"Chunk {i+1}/{total} failed: {e}")
                translated_parts.append(chunk)

            if progress_callback:
                progress_callback((i + 1) / total, f"ترجمة الجزء {i+1}/{total}")

        return "\n".join(translated_parts)

    def translate_batch_fast(self, texts: List[str], progress_callback: Optional[Callable] = None) -> List[str]:
        """
        ULTRA-FAST batch translation: combines multiple texts into mega-chunks,
        translates them in parallel using multi-threading.
        """
        if not texts:
            return []

        indexed_texts = []
        results = [""] * len(texts)

        for i, txt in enumerate(texts):
            if txt and txt.strip():
                indexed_texts.append((i, txt))
            else:
                results[i] = txt

        if not indexed_texts:
            return results

        # Build mega-chunks
        mega_chunks = []
        current_indices = []
        current_text = ""

        for idx, txt in indexed_texts:
            candidate = current_text + BATCH_SEPARATOR + txt if current_text else txt
            if len(candidate) <= CHUNK_SIZE:
                current_text = candidate
                current_indices.append(idx)
            else:
                if current_text:
                    mega_chunks.append((list(current_indices), current_text))
                if len(txt) > CHUNK_SIZE:
                    mega_chunks.append(([idx], txt))
                    current_text = ""
                    current_indices = []
                else:
                    current_text = txt
                    current_indices = [idx]

        if current_text:
            mega_chunks.append((list(current_indices), current_text))

        total_mega = len(mega_chunks)
        
        # USE MULTI-THREADING FOR SPEED
        from concurrent.futures import ThreadPoolExecutor
        
        def process_chunk(chunk_data):
            indices, combined = chunk_data
            try:
                translated = self._translator.translate(combined)
                if not translated: translated = combined
                
                if len(indices) == 1:
                    return [(indices[0], translated)]
                else:
                    parts = self._smart_split(translated, len(indices))
                    chunk_results = []
                    for j, part_idx in enumerate(indices):
                        val = parts[j].strip() if j < len(parts) else texts[part_idx]
                        chunk_results.append((part_idx, val))
                    return chunk_results
            except Exception as e:
                import logging
                logging.warning(f"Chunk failed: {e}")
                return [(idx, texts[idx]) for idx in indices]

        # Limit threads to avoid rate limiting (max 5)
        if self.logger:
            self.logger.log(f"🚀 معالجة {total_mega} حزمة ترجمة متوازية...")
        with ThreadPoolExecutor(max_workers=min(total_mega, 5)) as executor:
            batch_results = list(executor.map(process_chunk, mega_chunks))
            
            # Map back to final results
            completed = 0
            for chunk_res in batch_results:
                for idx, val in chunk_res:
                    results[idx] = val
                completed += 1
                if progress_callback:
                    progress_callback(completed / total_mega, f"مكتمل: {completed}/{total_mega}")

        return results

    def _smart_split(self, translated_text: str, expected_count: int) -> List[str]:
        """Split translated text back into parts, handling separator variations."""
        # Try exact separator first
        for sep in ["|||SPLIT|||", "|||split|||", "||| SPLIT |||", "|||Split|||",
                     "|||تقسيم|||", "|||انقسام|||", "SPLIT", "|||"]:
            parts = translated_text.split(sep)
            if len(parts) == expected_count:
                return parts

        # Fallback: split by double newline
        parts = translated_text.split("\n\n")
        if len(parts) >= expected_count:
            return parts[:expected_count]

        # Last resort: split by newline
        parts = translated_text.split("\n")
        if len(parts) >= expected_count:
            # Distribute lines evenly
            chunk_size = max(1, len(parts) // expected_count)
            result = []
            for i in range(expected_count):
                start = i * chunk_size
                end = start + chunk_size if i < expected_count - 1 else len(parts)
                result.append("\n".join(parts[start:end]))
            return result

        # Absolute fallback: return as single chunk
        return [translated_text]

    def _split_into_chunks(self, text: str) -> List[str]:
        """Split text into chunks at paragraph/sentence boundaries."""
        if len(text) <= CHUNK_SIZE:
            return [text]

        chunks = []
        current = ""

        for para in text.split("\n"):
            if len(current) + len(para) + 1 <= CHUNK_SIZE:
                current += para + "\n"
            else:
                if current:
                    chunks.append(current.strip())
                if len(para) > CHUNK_SIZE:
                    # Force-split very long paragraphs
                    for j in range(0, len(para), CHUNK_SIZE):
                        chunks.append(para[j:j + CHUNK_SIZE])
                    current = ""
                else:
                    current = para + "\n"

        if current.strip():
            chunks.append(current.strip())

        return chunks if chunks else [text]


# ──────────────────────────────────────────────
# Operation Logger
# ──────────────────────────────────────────────

class OperationLogger:
    """Logs translation operations with timestamps."""

    def __init__(self):
        self.entries: List[Dict[str, Any]] = []

    def log(self, message: str, level: str = "info"):
        self.entries.append({
            "time": datetime.now().strftime("%H:%M:%S"),
            "message": message,
            "level": level,
        })

    def get_entries(self) -> List[Dict[str, Any]]:
        return self.entries

    def clear(self):
        self.entries.clear()


# ──────────────────────────────────────────────
# Base Handler
# ──────────────────────────────────────────────

class BaseTranslator:
    """Base class for all file type translators."""

    SUPPORTED_EXTENSIONS: List[str] = []

    def __init__(self, service: TranslationService, logger: OperationLogger):
        self.service = service
        self.logger = logger

    def can_handle(self, extension: str) -> bool:
        return extension.lower() in self.SUPPORTED_EXTENSIONS

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        raise NotImplementedError

    def _error(self, msg):
        self.logger.log(f"❌ {msg}", "error")
        return {"success": False, "error": msg, "original_text": "", "translated_text": "",
                "output_bytes": b"", "output_filename": "", "output_mime": ""}


# ──────────────────────────────────────────────
# DOCX Translator (Optimized — batched)
# ──────────────────────────────────────────────

class DocxTranslator(BaseTranslator):
    """Translate .docx files while preserving formatting. Uses fast batch translation."""

    SUPPORTED_EXTENSIONS = [".docx"]

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        try:
            from docx import Document
        except ImportError:
            return self._error("مكتبة python-docx غير مثبتة.\nشغّل: pip install python-docx")

        self.logger.log(f"📄 بدء ترجمة ملف Word: {filename}")

        try:
            doc = Document(io.BytesIO(file_bytes))
            original_texts = []
            translatable_runs = []

            # Phase 1: Collect all text runs
            self.logger.log("📋 استخراج النصوص من المستند...")
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text and run.text.strip():
                        original_texts.append(run.text)
                        translatable_runs.append(run)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text and run.text.strip():
                                    original_texts.append(run.text)
                                    translatable_runs.append(run)

            if not original_texts:
                return self._error("لم يتم العثور على نصوص في الملف")

            total = len(original_texts)
            self.logger.log(f"📊 تم العثور على {total} نص قابل للترجمة")

            # Phase 2: FAST batch translation (combine → translate → split)
            self.logger.log("🚀 بدء الترجمة السريعة (دُفعة واحدة)...")
            translated_texts = self.service.translate_batch_fast(
                original_texts, progress_callback=progress_callback
            )

            # Phase 3: Replace text in-place (formatting preserved)
            self.logger.log("✏️ استبدال النصوص مع الحفاظ على التنسيق...")
            for run, trans in zip(translatable_runs, translated_texts):
                run.text = trans

            # Save
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            base_name = os.path.splitext(filename)[0]
            output_filename = f"translated_{base_name}.docx"
            self.logger.log(f"✅ تمت الترجمة بنجاح: {output_filename}", "success")

            return {
                "original_text": "\n".join(original_texts),
                "translated_text": "\n".join(translated_texts),
                "output_bytes": output.read(),
                "output_filename": output_filename,
                "output_mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "success": True, "error": None,
                "stats": {"total_texts": total}
            }

        except Exception as e:
            return self._error(f"خطأ في ترجمة ملف Word: {str(e)}")


# ──────────────────────────────────────────────
# Image Translator (OCR + Translation)
# Uses easyocr if available, otherwise pytesseract, otherwise Pillow-only
# ──────────────────────────────────────────────

class ImageTranslator(BaseTranslator):
    """Extract text from images via OCR, then translate."""

    SUPPORTED_EXTENSIONS = [".jpg", ".jpeg", ".png"]

    def _ocr_with_easyocr(self, image_bytes: bytes) -> Optional[List[str]]:
        """Try OCR with easyocr (best quality, but heavy)."""
        try:
            import easyocr
            import numpy as np
            from PIL import Image

            reader = easyocr.Reader(['en', 'ar'], gpu=False, verbose=False)
            image = Image.open(io.BytesIO(image_bytes))
            image_np = np.array(image)
            results = reader.readtext(image_np)
            if results:
                return [text for (_, text, _) in results]
        except ImportError:
            pass
        except Exception as e:
            logging.warning(f"easyocr failed: {e}")
        return None

    def _ocr_with_pytesseract(self, image_bytes: bytes) -> Optional[List[str]]:
        """Try OCR with pytesseract (lighter, needs Tesseract installed)."""
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(image, lang='eng+ara')
            if text and text.strip():
                return text.strip().split("\n")
        except ImportError:
            pass
        except Exception as e:
            logging.warning(f"pytesseract failed: {e}")
        return None

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:

        self.logger.log(f"🖼️ بدء معالجة الصورة: {filename}")

        if progress_callback:
            progress_callback(0.1, "جارٍ استخراج النص من الصورة...")

        # Try OCR methods in order of preference
        extracted_lines = self._ocr_with_easyocr(file_bytes)
        if extracted_lines:
            self.logger.log("✅ تم استخراج النص بـ EasyOCR")
        else:
            extracted_lines = self._ocr_with_pytesseract(file_bytes)
            if extracted_lines:
                self.logger.log("✅ تم استخراج النص بـ Tesseract")
            else:
                return self._error(
                    "لم يتم العثور على أداة OCR.\n"
                    "ثبّت إحدى المكتبات:\n"
                    "• pip install easyocr (أفضل جودة)\n"
                    "• pip install pytesseract (أخف وزناً، يحتاج Tesseract)"
                )

        # Filter empty lines
        extracted_lines = [l for l in extracted_lines if l and l.strip()]
        if not extracted_lines:
            return self._error("لم يتم العثور على نص في الصورة")

        original_text = "\n".join(extracted_lines)
        self.logger.log(f"📊 تم استخراج {len(extracted_lines)} سطر")

        if progress_callback:
            progress_callback(0.4, "جارٍ ترجمة النص المستخرج...")

        # Translate
        self.logger.log("🔄 بدء الترجمة...")

        def sub_progress(pct, msg):
            if progress_callback:
                progress_callback(0.4 + pct * 0.6, msg)

        translated_text = self.service.translate_text(original_text, progress_callback=sub_progress)

        base_name = os.path.splitext(filename)[0]
        output_filename = f"translated_{base_name}.txt"
        output_content = f"=== النص الأصلي (Original) ===\n\n{original_text}\n\n"
        output_content += f"=== النص المترجم (Translated) ===\n\n{translated_text}\n"

        self.logger.log(f"✅ تمت ترجمة نص الصورة: {output_filename}", "success")

        return {
            "original_text": original_text,
            "translated_text": translated_text,
            "output_bytes": output_content.encode("utf-8"),
            "output_filename": output_filename,
            "output_mime": "text/plain",
            "success": True, "error": None,
            "stats": {"lines_extracted": len(extracted_lines)}
        }


# ──────────────────────────────────────────────
# BDF Font File Translator
# ──────────────────────────────────────────────

class BdfTranslator(BaseTranslator):
    """Translate text properties in BDF font files."""

    SUPPORTED_EXTENSIONS = [".bdf"]

    TRANSLATABLE_KEYS = {
        "FAMILY_NAME", "FONT", "COMMENT", "COPYRIGHT",
        "NOTICE", "FOUNDRY", "FULL_NAME", "WEIGHT_NAME",
        "SLANT", "ADD_STYLE_NAME", "CHARSET_REGISTRY"
    }

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:

        self.logger.log(f"🔤 بدء ترجمة ملف BDF: {filename}")

        try:
            try:
                content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                content = file_bytes.decode("latin-1")

            lines = content.split("\n")
            total_lines = len(lines)
            self.logger.log(f"📊 الملف يحتوي على {total_lines} سطر")

            # Phase 1: Collect translatable texts
            translatable_items = []  # (line_index, key, value, is_quoted)
            for i, line in enumerate(lines):
                stripped = line.strip()
                for key in self.TRANSLATABLE_KEYS:
                    if stripped.startswith(key + " "):
                        value = stripped[len(key) + 1:].strip()
                        is_quoted = value.startswith('"') and value.endswith('"')
                        inner = value[1:-1] if is_quoted else value
                        if inner and not self._is_technical(inner):
                            translatable_items.append((i, key, inner, is_quoted))
                        break

            if not translatable_items:
                self.logger.log("ℹ️ لا توجد نصوص قابلة للترجمة في هذا الملف", "warning")
                return {
                    "original_text": "(لا يوجد نصوص قابلة للترجمة)",
                    "translated_text": "(لا يوجد ترجمة)",
                    "output_bytes": file_bytes,
                    "output_filename": f"translated_{os.path.splitext(filename)[0]}.bdf",
                    "output_mime": "application/octet-stream",
                    "success": True, "error": None,
                    "stats": {"total_lines": total_lines, "translated_properties": 0}
                }

            # Phase 2: Batch translate
            texts_to_translate = [item[2] for item in translatable_items]
            self.logger.log(f"🔄 ترجمة {len(texts_to_translate)} خاصية...")

            translated = self.service.translate_batch_fast(texts_to_translate, progress_callback)

            # Phase 3: Replace in content
            for (line_idx, key, original, is_quoted), trans in zip(translatable_items, translated):
                if is_quoted:
                    lines[line_idx] = f'{key} "{trans}"'
                else:
                    lines[line_idx] = f"{key} {trans}"

            output_content = "\n".join(lines)
            base_name = os.path.splitext(filename)[0]
            output_filename = f"translated_{base_name}.bdf"

            self.logger.log(f"✅ تمت ترجمة {len(translatable_items)} خاصية في BDF", "success")

            return {
                "original_text": "\n".join(texts_to_translate),
                "translated_text": "\n".join(translated),
                "output_bytes": output_content.encode("utf-8"),
                "output_filename": output_filename,
                "output_mime": "application/octet-stream",
                "success": True, "error": None,
                "stats": {"total_lines": total_lines, "translated_properties": len(translatable_items)}
            }

        except Exception as e:
            return self._error(f"خطأ في ترجمة BDF: {str(e)}")

    def _is_technical(self, text: str) -> bool:
        if re.match(r'^[0-9A-Fa-f\s\-_.]+$', text):
            return True
        if re.match(r'^ISO\d+', text):
            return True
        if len(text) <= 1:
            return True
        return False


# ──────────────────────────────────────────────
# PDF Translator
# ──────────────────────────────────────────────

class PdfTranslator(BaseTranslator):
    """Translate PDF files using pdfplumber for extraction."""

    SUPPORTED_EXTENSIONS = [".pdf"]

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        try:
            import pdfplumber
        except ImportError:
            return self._error("مكتبة pdfplumber غير مثبتة.\nشغّل: pip install pdfplumber")

        self.logger.log(f"📑 بدء ترجمة ملف PDF: {filename}")

        try:
            self.logger.log("📋 استخراج النصوص من PDF...")
            if progress_callback:
                progress_callback(0.05, "جارٍ استخراج النصوص...")

            original_text = ""
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                total_pages = len(pdf.pages)
                for idx, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        original_text += page_text + "\n\n"
                    if progress_callback:
                        progress_callback(0.05 + 0.15 * ((idx + 1) / total_pages),
                                          f"صفحة {idx+1}/{total_pages}")

            if not original_text.strip():
                return self._error("لم يتم العثور على نص في PDF. قد يكون عبارة عن صور.")

            self.logger.log(f"📊 تم استخراج {len(original_text)} حرف من {total_pages} صفحة")

            # Translate
            self.logger.log("🚀 بدء الترجمة...")

            def sub_progress(pct, msg):
                if progress_callback:
                    progress_callback(0.2 + pct * 0.8, msg)

            translated_text = self.service.translate_text(original_text, progress_callback=sub_progress)

            base_name = os.path.splitext(filename)[0]
            output_filename = f"translated_{base_name}.txt"

            self.logger.log(f"✅ تمت ترجمة PDF بنجاح", "success")

            return {
                "original_text": original_text,
                "translated_text": translated_text,
                "output_bytes": translated_text.encode("utf-8"),
                "output_filename": output_filename,
                "output_mime": "text/plain",
                "success": True, "error": None,
                "stats": {"total_pages": total_pages, "total_chars": len(original_text)}
            }

        except Exception as e:
            return self._error(f"خطأ في ترجمة PDF: {str(e)}")


# ──────────────────────────────────────────────
# Main File Translator (Orchestrator)
# ──────────────────────────────────────────────

class FileTranslator:
    """
    Main orchestrator: auto-detects file type → routes to correct handler.

    Usage:
        translator = FileTranslator(source_lang="auto", target_lang="ar")
        result = translator.translate(file_bytes, "resume.docx", progress_callback)
    """

    SUPPORTED_EXTENSIONS = [".docx", ".bdf", ".jpg", ".jpeg", ".png", ".pdf"]

    FILE_TYPE_INFO = {
        ".docx": {"icon": "📄", "name_ar": "ملف Word", "name_en": "Word Document"},
        ".bdf":  {"icon": "🔤", "name_ar": "ملف خط BDF", "name_en": "BDF Font File"},
        ".jpg":  {"icon": "🖼️", "name_ar": "صورة JPG", "name_en": "JPG Image"},
        ".jpeg": {"icon": "🖼️", "name_ar": "صورة JPEG", "name_en": "JPEG Image"},
        ".png":  {"icon": "🖼️", "name_ar": "صورة PNG", "name_en": "PNG Image"},
        ".pdf":  {"icon": "📑", "name_ar": "ملف PDF", "name_en": "PDF Document"},
    }

    def __init__(self, source_lang: str = "auto", target_lang: str = "ar"):
        self.logger = OperationLogger()
        self.service = TranslationService(source_lang=source_lang, target_lang=target_lang, logger=self.logger)
        self._handlers: List[BaseTranslator] = [
            DocxTranslator(self.service, self.logger),
            ImageTranslator(self.service, self.logger),
            BdfTranslator(self.service, self.logger),
            PdfTranslator(self.service, self.logger),
        ]

    def get_file_type(self, filename: str) -> Optional[str]:
        _, ext = os.path.splitext(filename)
        return ext.lower() if ext else None

    def get_handler(self, filename: str) -> Optional[BaseTranslator]:
        ext = self.get_file_type(filename)
        if not ext:
            return None
        for handler in self._handlers:
            if handler.can_handle(ext):
                return handler
        return None

    def is_supported(self, filename: str) -> bool:
        return self.get_file_type(filename) in self.SUPPORTED_EXTENSIONS

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        ext = self.get_file_type(filename)
        info = self.FILE_TYPE_INFO.get(ext, {})
        icon = info.get("icon", "📁")
        type_name = info.get("name_ar", ext)

        self.logger.log(f"📂 تم استلام ملف: {filename}")
        self.logger.log(f"{icon} نوع الملف: {type_name}")
        self.logger.log(f"📏 حجم الملف: {self._format_size(len(file_bytes))}")

        handler = self.get_handler(filename)
        if not handler:
            supported = ", ".join(self.SUPPORTED_EXTENSIONS)
            msg = f"نوع الملف '{ext}' غير مدعوم. الأنواع المدعومة: {supported}"
            self.logger.log(f"❌ {msg}", "error")
            return {"success": False, "error": msg}

        start_time = time.time()
        result = handler.translate(file_bytes, filename, progress_callback)
        elapsed = time.time() - start_time

        if result.get("success"):
            self.logger.log(f"⏱️ إجمالي وقت الترجمة: {elapsed:.1f} ثانية", "success")
        else:
            self.logger.log(f"⏱️ فشل بعد: {elapsed:.1f} ثانية", "error")

        return result

    def get_log(self) -> List[Dict]:
        return self.logger.get_entries()

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
=======
# src/core/file_translator.py
"""
Professional Multi-Format File Translation Engine (Optimized)
Supports: .docx, .bdf, .jpg/.png/.jpeg, .pdf
Uses: deep_translator (free, no API key)
Architecture: Extensible handler pattern for future formats (pptx, xlsx)
Author: Alsaeed Alwazzan

Performance Optimizations:
- Batch translation: combine texts with separator → single API call → split back
- Larger chunks (4500 chars) to minimize API roundtrips
- Lazy imports to avoid startup delay
- OCR is optional (easyocr or pytesseract, with graceful fallback)
"""

import io
import os
import re
import time
import logging
from datetime import datetime
from typing import Callable, Optional, Dict, List, Tuple, Any

# ──────────────────────────────────────────────
# Constants
# ──────────────────────────────────────────────
BATCH_SEPARATOR = "\n|||SPLIT|||\n"
CHUNK_SIZE = 4500  # Max chars per translation request


# ──────────────────────────────────────────────
# Translation Service (Free, No API Key)
# ──────────────────────────────────────────────

class TranslationService:
    """Fast, chunked translation using deep_translator (free Google Translate)."""

    SUPPORTED_LANGUAGES = {
        "ar": "العربية (Arabic)",
        "en": "English",
        "fr": "الفرنسية (French)",
        "es": "الإسبانية (Spanish)",
        "de": "الألمانية (German)",
        "tr": "التركية (Turkish)",
        "ur": "الأردية (Urdu)",
        "hi": "الهندية (Hindi)",
        "zh-CN": "الصينية (Chinese)",
        "ja": "اليابانية (Japanese)",
        "ko": "الكورية (Korean)",
        "ru": "الروسية (Russian)",
        "pt": "البرتغالية (Portuguese)",
        "it": "الإيطالية (Italian)",
        "id": "الإندونيسية (Indonesian)",
        "ms": "الملايوية (Malay)",
        "tl": "الفلبينية (Filipino)",
        "bn": "البنغالية (Bengali)",
        "th": "التايلاندية (Thai)",
        "vi": "الفيتنامية (Vietnamese)",
    }

    def __init__(self, source_lang: str = "auto", target_lang: str = "ar", logger=None):
        from deep_translator import GoogleTranslator
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.logger = logger
        self._translator = GoogleTranslator(source=source_lang, target=target_lang)

    def translate_text(self, text: str, progress_callback: Optional[Callable] = None) -> str:
        """Translate a single text string, with chunking for large content."""
        if not text or not text.strip():
            return ""

        # If small enough, translate directly
        if len(text) <= CHUNK_SIZE:
            try:
                result = self._translator.translate(text)
                if progress_callback:
                    progress_callback(1.0, "✅")
                return result if result else text
            except Exception as e:
                logging.warning(f"Translation failed: {e}")
                return text

        # Split large text into chunks at newline boundaries
        chunks = self._split_into_chunks(text)
        total = len(chunks)
        translated_parts = []

        for i, chunk in enumerate(chunks):
            try:
                result = self._translator.translate(chunk)
                translated_parts.append(result if result else chunk)
            except Exception as e:
                logging.warning(f"Chunk {i+1}/{total} failed: {e}")
                translated_parts.append(chunk)

            if progress_callback:
                progress_callback((i + 1) / total, f"ترجمة الجزء {i+1}/{total}")

        return "\n".join(translated_parts)

    def translate_batch_fast(self, texts: List[str], progress_callback: Optional[Callable] = None) -> List[str]:
        """
        ULTRA-FAST batch translation: combines multiple texts into mega-chunks,
        translates them in parallel using multi-threading.
        """
        if not texts:
            return []

        indexed_texts = []
        results = [""] * len(texts)

        for i, txt in enumerate(texts):
            if txt and txt.strip():
                indexed_texts.append((i, txt))
            else:
                results[i] = txt

        if not indexed_texts:
            return results

        # Build mega-chunks
        mega_chunks = []
        current_indices = []
        current_text = ""

        for idx, txt in indexed_texts:
            candidate = current_text + BATCH_SEPARATOR + txt if current_text else txt
            if len(candidate) <= CHUNK_SIZE:
                current_text = candidate
                current_indices.append(idx)
            else:
                if current_text:
                    mega_chunks.append((list(current_indices), current_text))
                if len(txt) > CHUNK_SIZE:
                    mega_chunks.append(([idx], txt))
                    current_text = ""
                    current_indices = []
                else:
                    current_text = txt
                    current_indices = [idx]

        if current_text:
            mega_chunks.append((list(current_indices), current_text))

        total_mega = len(mega_chunks)
        
        # USE MULTI-THREADING FOR SPEED
        from concurrent.futures import ThreadPoolExecutor
        
        def process_chunk(chunk_data):
            indices, combined = chunk_data
            try:
                translated = self._translator.translate(combined)
                if not translated: translated = combined
                
                if len(indices) == 1:
                    return [(indices[0], translated)]
                else:
                    parts = self._smart_split(translated, len(indices))
                    chunk_results = []
                    for j, part_idx in enumerate(indices):
                        val = parts[j].strip() if j < len(parts) else texts[part_idx]
                        chunk_results.append((part_idx, val))
                    return chunk_results
            except Exception as e:
                import logging
                logging.warning(f"Chunk failed: {e}")
                return [(idx, texts[idx]) for idx in indices]

        # Limit threads to avoid rate limiting (max 5)
        if self.logger:
            self.logger.log(f"🚀 معالجة {total_mega} حزمة ترجمة متوازية...")
        with ThreadPoolExecutor(max_workers=min(total_mega, 5)) as executor:
            batch_results = list(executor.map(process_chunk, mega_chunks))
            
            # Map back to final results
            completed = 0
            for chunk_res in batch_results:
                for idx, val in chunk_res:
                    results[idx] = val
                completed += 1
                if progress_callback:
                    progress_callback(completed / total_mega, f"مكتمل: {completed}/{total_mega}")

        return results

    def _smart_split(self, translated_text: str, expected_count: int) -> List[str]:
        """Split translated text back into parts, handling separator variations."""
        # Try exact separator first
        for sep in ["|||SPLIT|||", "|||split|||", "||| SPLIT |||", "|||Split|||",
                     "|||تقسيم|||", "|||انقسام|||", "SPLIT", "|||"]:
            parts = translated_text.split(sep)
            if len(parts) == expected_count:
                return parts

        # Fallback: split by double newline
        parts = translated_text.split("\n\n")
        if len(parts) >= expected_count:
            return parts[:expected_count]

        # Last resort: split by newline
        parts = translated_text.split("\n")
        if len(parts) >= expected_count:
            # Distribute lines evenly
            chunk_size = max(1, len(parts) // expected_count)
            result = []
            for i in range(expected_count):
                start = i * chunk_size
                end = start + chunk_size if i < expected_count - 1 else len(parts)
                result.append("\n".join(parts[start:end]))
            return result

        # Absolute fallback: return as single chunk
        return [translated_text]

    def _split_into_chunks(self, text: str) -> List[str]:
        """Split text into chunks at paragraph/sentence boundaries."""
        if len(text) <= CHUNK_SIZE:
            return [text]

        chunks = []
        current = ""

        for para in text.split("\n"):
            if len(current) + len(para) + 1 <= CHUNK_SIZE:
                current += para + "\n"
            else:
                if current:
                    chunks.append(current.strip())
                if len(para) > CHUNK_SIZE:
                    # Force-split very long paragraphs
                    for j in range(0, len(para), CHUNK_SIZE):
                        chunks.append(para[j:j + CHUNK_SIZE])
                    current = ""
                else:
                    current = para + "\n"

        if current.strip():
            chunks.append(current.strip())

        return chunks if chunks else [text]


# ──────────────────────────────────────────────
# Operation Logger
# ──────────────────────────────────────────────

class OperationLogger:
    """Logs translation operations with timestamps."""

    def __init__(self):
        self.entries: List[Dict[str, Any]] = []

    def log(self, message: str, level: str = "info"):
        self.entries.append({
            "time": datetime.now().strftime("%H:%M:%S"),
            "message": message,
            "level": level,
        })

    def get_entries(self) -> List[Dict[str, Any]]:
        return self.entries

    def clear(self):
        self.entries.clear()


# ──────────────────────────────────────────────
# Base Handler
# ──────────────────────────────────────────────

class BaseTranslator:
    """Base class for all file type translators."""

    SUPPORTED_EXTENSIONS: List[str] = []

    def __init__(self, service: TranslationService, logger: OperationLogger):
        self.service = service
        self.logger = logger

    def can_handle(self, extension: str) -> bool:
        return extension.lower() in self.SUPPORTED_EXTENSIONS

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        raise NotImplementedError

    def _error(self, msg):
        self.logger.log(f"❌ {msg}", "error")
        return {"success": False, "error": msg, "original_text": "", "translated_text": "",
                "output_bytes": b"", "output_filename": "", "output_mime": ""}


# ──────────────────────────────────────────────
# DOCX Translator (Optimized — batched)
# ──────────────────────────────────────────────

class DocxTranslator(BaseTranslator):
    """Translate .docx files while preserving formatting. Uses fast batch translation."""

    SUPPORTED_EXTENSIONS = [".docx"]

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        try:
            from docx import Document
        except ImportError:
            return self._error("مكتبة python-docx غير مثبتة.\nشغّل: pip install python-docx")

        self.logger.log(f"📄 بدء ترجمة ملف Word: {filename}")

        try:
            doc = Document(io.BytesIO(file_bytes))
            original_texts = []
            translatable_runs = []

            # Phase 1: Collect all text runs
            self.logger.log("📋 استخراج النصوص من المستند...")
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text and run.text.strip():
                        original_texts.append(run.text)
                        translatable_runs.append(run)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text and run.text.strip():
                                    original_texts.append(run.text)
                                    translatable_runs.append(run)

            if not original_texts:
                return self._error("لم يتم العثور على نصوص في الملف")

            total = len(original_texts)
            self.logger.log(f"📊 تم العثور على {total} نص قابل للترجمة")

            # Phase 2: FAST batch translation (combine → translate → split)
            self.logger.log("🚀 بدء الترجمة السريعة (دُفعة واحدة)...")
            translated_texts = self.service.translate_batch_fast(
                original_texts, progress_callback=progress_callback
            )

            # Phase 3: Replace text in-place (formatting preserved)
            self.logger.log("✏️ استبدال النصوص مع الحفاظ على التنسيق...")
            for run, trans in zip(translatable_runs, translated_texts):
                run.text = trans

            # Save
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            base_name = os.path.splitext(filename)[0]
            output_filename = f"translated_{base_name}.docx"
            self.logger.log(f"✅ تمت الترجمة بنجاح: {output_filename}", "success")

            return {
                "original_text": "\n".join(original_texts),
                "translated_text": "\n".join(translated_texts),
                "output_bytes": output.read(),
                "output_filename": output_filename,
                "output_mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "success": True, "error": None,
                "stats": {"total_texts": total}
            }

        except Exception as e:
            return self._error(f"خطأ في ترجمة ملف Word: {str(e)}")


# ──────────────────────────────────────────────
# Image Translator (OCR + Translation)
# Uses easyocr if available, otherwise pytesseract, otherwise Pillow-only
# ──────────────────────────────────────────────

class ImageTranslator(BaseTranslator):
    """Extract text from images via OCR, then translate."""

    SUPPORTED_EXTENSIONS = [".jpg", ".jpeg", ".png"]

    def _ocr_with_easyocr(self, image_bytes: bytes) -> Optional[List[str]]:
        """Try OCR with easyocr (best quality, but heavy)."""
        try:
            import easyocr
            import numpy as np
            from PIL import Image

            reader = easyocr.Reader(['en', 'ar'], gpu=False, verbose=False)
            image = Image.open(io.BytesIO(image_bytes))
            image_np = np.array(image)
            results = reader.readtext(image_np)
            if results:
                return [text for (_, text, _) in results]
        except ImportError:
            pass
        except Exception as e:
            logging.warning(f"easyocr failed: {e}")
        return None

    def _ocr_with_pytesseract(self, image_bytes: bytes) -> Optional[List[str]]:
        """Try OCR with pytesseract (lighter, needs Tesseract installed)."""
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(image, lang='eng+ara')
            if text and text.strip():
                return text.strip().split("\n")
        except ImportError:
            pass
        except Exception as e:
            logging.warning(f"pytesseract failed: {e}")
        return None

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:

        self.logger.log(f"🖼️ بدء معالجة الصورة: {filename}")

        if progress_callback:
            progress_callback(0.1, "جارٍ استخراج النص من الصورة...")

        # Try OCR methods in order of preference
        extracted_lines = self._ocr_with_easyocr(file_bytes)
        if extracted_lines:
            self.logger.log("✅ تم استخراج النص بـ EasyOCR")
        else:
            extracted_lines = self._ocr_with_pytesseract(file_bytes)
            if extracted_lines:
                self.logger.log("✅ تم استخراج النص بـ Tesseract")
            else:
                return self._error(
                    "لم يتم العثور على أداة OCR.\n"
                    "ثبّت إحدى المكتبات:\n"
                    "• pip install easyocr (أفضل جودة)\n"
                    "• pip install pytesseract (أخف وزناً، يحتاج Tesseract)"
                )

        # Filter empty lines
        extracted_lines = [l for l in extracted_lines if l and l.strip()]
        if not extracted_lines:
            return self._error("لم يتم العثور على نص في الصورة")

        original_text = "\n".join(extracted_lines)
        self.logger.log(f"📊 تم استخراج {len(extracted_lines)} سطر")

        if progress_callback:
            progress_callback(0.4, "جارٍ ترجمة النص المستخرج...")

        # Translate
        self.logger.log("🔄 بدء الترجمة...")

        def sub_progress(pct, msg):
            if progress_callback:
                progress_callback(0.4 + pct * 0.6, msg)

        translated_text = self.service.translate_text(original_text, progress_callback=sub_progress)

        base_name = os.path.splitext(filename)[0]
        output_filename = f"translated_{base_name}.txt"
        output_content = f"=== النص الأصلي (Original) ===\n\n{original_text}\n\n"
        output_content += f"=== النص المترجم (Translated) ===\n\n{translated_text}\n"

        self.logger.log(f"✅ تمت ترجمة نص الصورة: {output_filename}", "success")

        return {
            "original_text": original_text,
            "translated_text": translated_text,
            "output_bytes": output_content.encode("utf-8"),
            "output_filename": output_filename,
            "output_mime": "text/plain",
            "success": True, "error": None,
            "stats": {"lines_extracted": len(extracted_lines)}
        }


# ──────────────────────────────────────────────
# BDF Font File Translator
# ──────────────────────────────────────────────

class BdfTranslator(BaseTranslator):
    """Translate text properties in BDF font files."""

    SUPPORTED_EXTENSIONS = [".bdf"]

    TRANSLATABLE_KEYS = {
        "FAMILY_NAME", "FONT", "COMMENT", "COPYRIGHT",
        "NOTICE", "FOUNDRY", "FULL_NAME", "WEIGHT_NAME",
        "SLANT", "ADD_STYLE_NAME", "CHARSET_REGISTRY"
    }

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:

        self.logger.log(f"🔤 بدء ترجمة ملف BDF: {filename}")

        try:
            try:
                content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                content = file_bytes.decode("latin-1")

            lines = content.split("\n")
            total_lines = len(lines)
            self.logger.log(f"📊 الملف يحتوي على {total_lines} سطر")

            # Phase 1: Collect translatable texts
            translatable_items = []  # (line_index, key, value, is_quoted)
            for i, line in enumerate(lines):
                stripped = line.strip()
                for key in self.TRANSLATABLE_KEYS:
                    if stripped.startswith(key + " "):
                        value = stripped[len(key) + 1:].strip()
                        is_quoted = value.startswith('"') and value.endswith('"')
                        inner = value[1:-1] if is_quoted else value
                        if inner and not self._is_technical(inner):
                            translatable_items.append((i, key, inner, is_quoted))
                        break

            if not translatable_items:
                self.logger.log("ℹ️ لا توجد نصوص قابلة للترجمة في هذا الملف", "warning")
                return {
                    "original_text": "(لا يوجد نصوص قابلة للترجمة)",
                    "translated_text": "(لا يوجد ترجمة)",
                    "output_bytes": file_bytes,
                    "output_filename": f"translated_{os.path.splitext(filename)[0]}.bdf",
                    "output_mime": "application/octet-stream",
                    "success": True, "error": None,
                    "stats": {"total_lines": total_lines, "translated_properties": 0}
                }

            # Phase 2: Batch translate
            texts_to_translate = [item[2] for item in translatable_items]
            self.logger.log(f"🔄 ترجمة {len(texts_to_translate)} خاصية...")

            translated = self.service.translate_batch_fast(texts_to_translate, progress_callback)

            # Phase 3: Replace in content
            for (line_idx, key, original, is_quoted), trans in zip(translatable_items, translated):
                if is_quoted:
                    lines[line_idx] = f'{key} "{trans}"'
                else:
                    lines[line_idx] = f"{key} {trans}"

            output_content = "\n".join(lines)
            base_name = os.path.splitext(filename)[0]
            output_filename = f"translated_{base_name}.bdf"

            self.logger.log(f"✅ تمت ترجمة {len(translatable_items)} خاصية في BDF", "success")

            return {
                "original_text": "\n".join(texts_to_translate),
                "translated_text": "\n".join(translated),
                "output_bytes": output_content.encode("utf-8"),
                "output_filename": output_filename,
                "output_mime": "application/octet-stream",
                "success": True, "error": None,
                "stats": {"total_lines": total_lines, "translated_properties": len(translatable_items)}
            }

        except Exception as e:
            return self._error(f"خطأ في ترجمة BDF: {str(e)}")

    def _is_technical(self, text: str) -> bool:
        if re.match(r'^[0-9A-Fa-f\s\-_.]+$', text):
            return True
        if re.match(r'^ISO\d+', text):
            return True
        if len(text) <= 1:
            return True
        return False


# ──────────────────────────────────────────────
# PDF Translator
# ──────────────────────────────────────────────

class PdfTranslator(BaseTranslator):
    """Translate PDF files using pdfplumber for extraction."""

    SUPPORTED_EXTENSIONS = [".pdf"]

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        try:
            import pdfplumber
        except ImportError:
            return self._error("مكتبة pdfplumber غير مثبتة.\nشغّل: pip install pdfplumber")

        self.logger.log(f"📑 بدء ترجمة ملف PDF: {filename}")

        try:
            self.logger.log("📋 استخراج النصوص من PDF...")
            if progress_callback:
                progress_callback(0.05, "جارٍ استخراج النصوص...")

            original_text = ""
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                total_pages = len(pdf.pages)
                for idx, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        original_text += page_text + "\n\n"
                    if progress_callback:
                        progress_callback(0.05 + 0.15 * ((idx + 1) / total_pages),
                                          f"صفحة {idx+1}/{total_pages}")

            if not original_text.strip():
                return self._error("لم يتم العثور على نص في PDF. قد يكون عبارة عن صور.")

            self.logger.log(f"📊 تم استخراج {len(original_text)} حرف من {total_pages} صفحة")

            # Translate
            self.logger.log("🚀 بدء الترجمة...")

            def sub_progress(pct, msg):
                if progress_callback:
                    progress_callback(0.2 + pct * 0.8, msg)

            translated_text = self.service.translate_text(original_text, progress_callback=sub_progress)

            base_name = os.path.splitext(filename)[0]
            output_filename = f"translated_{base_name}.txt"

            self.logger.log(f"✅ تمت ترجمة PDF بنجاح", "success")

            return {
                "original_text": original_text,
                "translated_text": translated_text,
                "output_bytes": translated_text.encode("utf-8"),
                "output_filename": output_filename,
                "output_mime": "text/plain",
                "success": True, "error": None,
                "stats": {"total_pages": total_pages, "total_chars": len(original_text)}
            }

        except Exception as e:
            return self._error(f"خطأ في ترجمة PDF: {str(e)}")


# ──────────────────────────────────────────────
# Main File Translator (Orchestrator)
# ──────────────────────────────────────────────

class FileTranslator:
    """
    Main orchestrator: auto-detects file type → routes to correct handler.

    Usage:
        translator = FileTranslator(source_lang="auto", target_lang="ar")
        result = translator.translate(file_bytes, "resume.docx", progress_callback)
    """

    SUPPORTED_EXTENSIONS = [".docx", ".bdf", ".jpg", ".jpeg", ".png", ".pdf"]

    FILE_TYPE_INFO = {
        ".docx": {"icon": "📄", "name_ar": "ملف Word", "name_en": "Word Document"},
        ".bdf":  {"icon": "🔤", "name_ar": "ملف خط BDF", "name_en": "BDF Font File"},
        ".jpg":  {"icon": "🖼️", "name_ar": "صورة JPG", "name_en": "JPG Image"},
        ".jpeg": {"icon": "🖼️", "name_ar": "صورة JPEG", "name_en": "JPEG Image"},
        ".png":  {"icon": "🖼️", "name_ar": "صورة PNG", "name_en": "PNG Image"},
        ".pdf":  {"icon": "📑", "name_ar": "ملف PDF", "name_en": "PDF Document"},
    }

    def __init__(self, source_lang: str = "auto", target_lang: str = "ar"):
        self.logger = OperationLogger()
        self.service = TranslationService(source_lang=source_lang, target_lang=target_lang, logger=self.logger)
        self._handlers: List[BaseTranslator] = [
            DocxTranslator(self.service, self.logger),
            ImageTranslator(self.service, self.logger),
            BdfTranslator(self.service, self.logger),
            PdfTranslator(self.service, self.logger),
        ]

    def get_file_type(self, filename: str) -> Optional[str]:
        _, ext = os.path.splitext(filename)
        return ext.lower() if ext else None

    def get_handler(self, filename: str) -> Optional[BaseTranslator]:
        ext = self.get_file_type(filename)
        if not ext:
            return None
        for handler in self._handlers:
            if handler.can_handle(ext):
                return handler
        return None

    def is_supported(self, filename: str) -> bool:
        return self.get_file_type(filename) in self.SUPPORTED_EXTENSIONS

    def translate(self, file_bytes: bytes, filename: str,
                  progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        ext = self.get_file_type(filename)
        info = self.FILE_TYPE_INFO.get(ext, {})
        icon = info.get("icon", "📁")
        type_name = info.get("name_ar", ext)

        self.logger.log(f"📂 تم استلام ملف: {filename}")
        self.logger.log(f"{icon} نوع الملف: {type_name}")
        self.logger.log(f"📏 حجم الملف: {self._format_size(len(file_bytes))}")

        handler = self.get_handler(filename)
        if not handler:
            supported = ", ".join(self.SUPPORTED_EXTENSIONS)
            msg = f"نوع الملف '{ext}' غير مدعوم. الأنواع المدعومة: {supported}"
            self.logger.log(f"❌ {msg}", "error")
            return {"success": False, "error": msg}

        start_time = time.time()
        result = handler.translate(file_bytes, filename, progress_callback)
        elapsed = time.time() - start_time

        if result.get("success"):
            self.logger.log(f"⏱️ إجمالي وقت الترجمة: {elapsed:.1f} ثانية", "success")
        else:
            self.logger.log(f"⏱️ فشل بعد: {elapsed:.1f} ثانية", "error")

        return result

    def get_log(self) -> List[Dict]:
        return self.logger.get_entries()

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
>>>>>>> 947f1af (update)
